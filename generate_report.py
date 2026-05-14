"""
generate_report.py -- Generates the GesturePlay project report as a .docx file.
Run:  pip install python-docx  (if not installed)
Then: python generate_report.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

OUTPUT = os.path.join(os.path.dirname(os.path.abspath(__file__)), "GesturePlay_Report.docx")


def set_cell_shading(cell, color_hex):
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn("w:shd"))
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.3
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table


def build_report():
    doc = Document()

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)

    # ==========================================================
    # TITLE PAGE
    # ==========================================================
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GesturePlay")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    run.font.name = "Calibri"

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Real-Time Hand Gesture-Controlled Media Player")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    run.font.name = "Calibri"

    doc.add_paragraph()

    course_info = doc.add_paragraph()
    course_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = course_info.add_run("Course: Digital Image Processing")
    run.font.size = Pt(12)
    run.font.name = "Calibri"

    doc.add_paragraph()

    tech = doc.add_paragraph()
    tech.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = tech.add_run("Python  •  OpenCV  •  MediaPipe  •  Flask  •  PyAutoGUI")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
    run.font.name = "Calibri"

    doc.add_page_break()

    # ==========================================================
    # 1. INTRODUCTION
    # ==========================================================
    add_heading(doc, "1. Introduction")

    add_para(doc,
        "GesturePlay is a desktop application that lets you control media playback "
        "using hand gestures detected through your laptop's webcam. The idea is "
        "straightforward: instead of reaching for the keyboard every time you want "
        "to pause a video or skip ahead, you just show a specific hand gesture to "
        "the camera and the system does the rest."
    )
    add_para(doc,
        "The entire system runs locally on your machine. There are no cloud APIs, "
        "no external servers, and no internet connection required once the "
        "dependencies are installed. A lightweight Flask web server provides a "
        "control dashboard where you can monitor the live camera feed, see which "
        "gesture is being detected, and remap gestures to different actions."
    )
    add_para(doc,
        "This project was built as part of the Digital Image Processing course "
        "to demonstrate practical applications of real-time computer vision and "
        "image processing techniques."
    )

    # ==========================================================
    # 2. PROBLEM STATEMENT
    # ==========================================================
    add_heading(doc, "2. Problem Statement")

    add_para(doc,
        "Most of us watch videos or movies on our laptops while sitting at a "
        "distance — maybe on a bed, across a desk, or just leaning back in a chair. "
        "Every time we need to pause, skip, or adjust the volume, we have to "
        "physically reach for the keyboard. It breaks the flow and gets annoying "
        "pretty quickly."
    )
    add_para(doc,
        "There are wireless remotes and voice assistants that can help, but they "
        "require extra hardware or subscriptions. Most laptops already have a "
        "built-in webcam sitting right there doing nothing. GesturePlay takes "
        "advantage of that existing hardware to provide a touchless media remote "
        "without any additional cost."
    )

    # ==========================================================
    # 3. OBJECTIVES
    # ==========================================================
    add_heading(doc, "3. Objectives")

    objectives = [
        "Detect and classify hand gestures in real time using the laptop webcam",
        "Map each gesture to a specific media control action (play, pause, seek, speed, mute)",
        "Build a stable detection pipeline that does not misfire on transitional hand positions",
        "Provide a web-based dashboard for live monitoring and configuration",
        "Work with YouTube in Chrome and Windows Media Player without browser extensions",
    ]
    for obj in objectives:
        p = doc.add_paragraph(obj, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"

    # ==========================================================
    # 4. SYSTEM ARCHITECTURE
    # ==========================================================
    add_heading(doc, "4. System Architecture")

    add_para(doc,
        "The system is organized as a simple pipeline. Each stage feeds directly "
        "into the next one, and the whole thing runs in a single Python process "
        "with two threads — one for the detection loop and one for the Flask web server."
    )
    add_para(doc, "The pipeline is as follows:", bold=True)

    pipeline_steps = [
        ("Webcam Capture", "OpenCV grabs frames from the webcam at 640×480, 30 FPS. "
         "The camera buffer is set to 1 to always get the freshest frame."),
        ("Hand Detection", "Each frame is passed to MediaPipe Hands, which returns "
         "21 landmark points for each detected hand."),
        ("Finger Counting", "The system checks which fingertips are above their "
         "corresponding knuckle joints to determine how many fingers are raised."),
        ("Gesture Classification", "The finger count maps directly to a gesture "
         "name: 0 = fist, 1 = one finger, 2 = peace, 3 = three fingers, 4/5 = open palm."),
        ("Stability Filter", "An 8-frame rolling buffer ensures a gesture must be "
         "consistent across at least 65% of recent frames before it is accepted."),
        ("Action Execution", "Once a gesture is confirmed and held long enough, "
         "PyAutoGUI sends the corresponding keyboard shortcut to the active window."),
    ]
    for step_name, step_desc in pipeline_steps:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"{step_name}: ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run2 = p.add_run(step_desc)
        run2.font.size = Pt(11)
        run2.font.name = "Calibri"

    add_para(doc,
        "A Flask web server runs alongside the detection thread. It streams the "
        "annotated camera feed as MJPEG, serves the web dashboard, and exposes REST "
        "API endpoints for gesture configuration and status polling."
    )

    # ==========================================================
    # 5. TECHNOLOGY STACK
    # ==========================================================
    add_heading(doc, "5. Technology Stack")

    add_table(doc,
        ["Component", "Library", "Purpose"],
        [
            ["Hand Tracking", "MediaPipe 0.10.x", "21-point hand landmark detection in real time"],
            ["Computer Vision", "OpenCV 4.10.x", "Webcam capture, frame flipping, annotation drawing"],
            ["Keyboard Control", "PyAutoGUI 0.9.x", "Sending keyboard shortcuts to the active window"],
            ["Window Detection", "PyGetWindow", "Checking which application is currently focused"],
            ["Web Server", "Flask 3.1.x", "Serving the dashboard UI and REST API"],
            ["Frontend", "HTML / CSS / JS", "Dashboard interface with live status updates"],
        ],
    )

    # ==========================================================
    # 6. GESTURE DETECTION METHODOLOGY
    # ==========================================================
    add_heading(doc, "6. Gesture Detection Methodology")

    # 6.1
    add_heading(doc, "6.1 MediaPipe Hand Landmarks", level=2)
    add_para(doc,
        "MediaPipe Hands is a pre-trained machine learning pipeline from Google "
        "that detects hands in an image and returns 21 3D landmark points. Each "
        "landmark has an x, y, and z coordinate. For our purposes, we only use "
        "the 2D pixel positions (x and y) since depth information is not needed "
        "for finger state detection."
    )
    add_para(doc,
        "The key landmarks used are the fingertip points (IDs 4, 8, 12, 16, 20) "
        "and their corresponding PIP joint points (IDs 6, 10, 14, 18). The thumb "
        "uses a horizontal distance check instead of a vertical one because it "
        "extends sideways rather than upward."
    )

    # 6.2
    add_heading(doc, "6.2 Finger State Detection", level=2)
    add_para(doc,
        "A finger is considered raised if its tip's Y-coordinate is less than "
        "its PIP joint's Y-coordinate (since Y increases downward in image space). "
        "In other words, if the tip is physically above the knuckle on screen, "
        "the finger is counted as extended."
    )
    add_para(doc,
        "The thumb requires different logic because it moves horizontally. "
        "We measure the horizontal distance between the thumb tip (landmark 4) "
        "and the thumb IP joint (landmark 2), normalized against the overall "
        "hand size. If this distance exceeds 40% of the hand size, the thumb "
        "is considered extended."
    )

    # 6.3
    add_heading(doc, "6.3 Gesture Classification", level=2)
    add_para(doc,
        "Gestures are classified by the total count of raised fingers. "
        "The mapping is kept deliberately simple to maximize reliability:"
    )
    add_table(doc,
        ["Finger Count", "Gesture Name", "Default Action"],
        [
            ["0", "Fist", "Play / Pause Toggle"],
            ["1", "One Finger", "Skip Backward 10s"],
            ["2", "Peace Sign", "Speed 2× (hold)"],
            ["3", "Three Fingers", "Skip Forward 10s"],
            ["4 or 5", "Open Palm", "Mute / Unmute"],
        ],
    )

    # 6.4
    add_heading(doc, "6.4 Stability Buffer", level=2)
    add_para(doc,
        "Raw per-frame detection is noisy. A hand moving between gestures, "
        "partial occlusion, or a brief lighting change can cause single-frame "
        "misclassifications. To handle this, the system maintains a rolling "
        "buffer of the last 8 raw detections using a Python deque."
    )
    add_para(doc,
        "A gesture is only considered stable when it appears in at least 65% "
        "of the buffer (roughly 5 out of 8 frames). This effectively filters "
        "out brief glitches without adding noticeable latency."
    )

    # 6.5
    add_heading(doc, "6.5 Neutral-Gap Gating", level=2)
    add_para(doc,
        "One of the trickiest problems we ran into was false triggers during "
        "hand transitions. For example, when you move from a fist (0 fingers) "
        "to an open palm (5 fingers), your hand briefly passes through 1, 2, "
        "3, and 4 fingers along the way. Without any safeguard, each of those "
        "intermediate states could fire an action."
    )
    add_para(doc,
        "The solution is neutral-gap gating. After any gesture triggers an "
        "action, the system locks out all further triggers until the hand is "
        "completely absent from the frame for at least 0.35 seconds. This "
        "forces the user to briefly lower their hand between gestures, which "
        "naturally prevents transition-state false positives."
    )

    # 6.6
    add_heading(doc, "6.6 Per-Gesture Hold Thresholds", level=2)
    add_para(doc,
        "Different gestures require different minimum hold durations before "
        "they fire. Destructive actions like play/pause need a longer hold "
        "(0.60 seconds) to prevent accidental triggers, while the speed-hold "
        "gesture responds faster (0.25 seconds) since it needs to feel "
        "responsive."
    )
    add_table(doc,
        ["Gesture", "Hold Duration", "Reason"],
        [
            ["Fist", "0.60 s", "Play/pause is high-impact; prevent misfires"],
            ["One Finger", "0.30 s", "Moderate — skip backward"],
            ["Peace Sign", "0.25 s", "Fast for hold-based speed control"],
            ["Three Fingers", "0.30 s", "Moderate — skip forward"],
            ["Open Palm", "0.40 s", "Mute is moderately impactful"],
        ],
    )

    # ==========================================================
    # 7. SPEED CONTROL
    # ==========================================================
    add_heading(doc, "7. Speed Control Mechanism")

    add_para(doc,
        "YouTube has a built-in feature where holding down the spacebar plays "
        "the video at 2× speed. GesturePlay uses this by pressing and holding "
        "the spacebar (via PyAutoGUI's keyDown) while the peace gesture is "
        "detected, and releasing it (keyUp) when the gesture changes. This "
        "gives smooth, continuous speed control that is directly tied to how "
        "long you hold the gesture."
    )

    # ==========================================================
    # 8. WINDOW FILTERING
    # ==========================================================
    add_heading(doc, "8. Window Filtering")

    add_para(doc,
        "Since PyAutoGUI sends keyboard shortcuts to whichever window is "
        "currently focused, gestures could accidentally affect other "
        "applications. For example, raising your hand while browsing a "
        "website could cause the page to scroll down."
    )
    add_para(doc,
        "To prevent this, the system includes an optional window filter. "
        "When enabled, it checks the title of the currently focused window "
        "using the PyGetWindow library. Keystrokes are only sent if the "
        "window title contains a recognized keyword (e.g. 'YouTube', 'VLC', "
        "'Netflix'). The list of allowed keywords is fully configurable."
    )

    # ==========================================================
    # 9. WEB DASHBOARD
    # ==========================================================
    add_heading(doc, "9. Web Dashboard")

    add_para(doc,
        "The Flask web application provides a real-time control panel with "
        "the following features:"
    )
    features = [
        ("Live Camera Feed", "An MJPEG stream of the annotated webcam view "
         "showing hand landmarks, the detected gesture, finger count, and FPS."),
        ("Gesture Readout", "A display panel showing the current gesture with "
         "its emoji and five indicator dots for the finger count."),
        ("Gesture Map Panel", "Dropdown selectors that let you remap any "
         "gesture to any available action without restarting the server."),
        ("Action Log", "A scrollable list of recently triggered actions with "
         "timestamps, updated in real time."),
        ("Start / Stop Toggle", "A single button to enable or disable "
         "gesture detection."),
    ]
    for fname, fdesc in features:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{fname}: ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run2 = p.add_run(fdesc)
        run2.font.size = Pt(11)
        run2.font.name = "Calibri"

    add_para(doc,
        "The dashboard uses a dark gray and silver color scheme, polls the "
        "status API every 400 milliseconds, and is fully responsive on "
        "mobile screens."
    )

    # ==========================================================
    # 10. PERFORMANCE OPTIMIZATIONS
    # ==========================================================
    add_heading(doc, "10. Performance Optimizations")

    add_table(doc,
        ["Optimization", "How It Works"],
        [
            ["Lightweight ML model", "MediaPipe model_complexity=0 (lite variant) for faster inference"],
            ["Camera buffer = 1", "CAP_PROP_BUFFERSIZE set to 1 so every read() gets the freshest frame"],
            ["Zero-wait frame sharing", "A deque(maxlen=1) passes frames between threads without locks"],
            ["JPEG quality 75", "Balances visual quality with fast encoding for the MJPEG stream"],
            ["640×480 resolution", "Sufficient for hand detection while keeping processing load low"],
        ],
    )

    # ==========================================================
    # 11. FILE STRUCTURE
    # ==========================================================
    add_heading(doc, "11. Project File Structure")

    add_table(doc,
        ["File", "Description"],
        [
            ["app.py", "Flask server, detection loop, MJPEG streaming, REST API"],
            ["gesture_detector.py", "MediaPipe hand tracking, finger counting, stability buffer"],
            ["gesture_controller.py", "Action execution via PyAutoGUI, window filtering, action log"],
            ["config_manager.py", "JSON configuration loading, saving, and runtime updates"],
            ["gesture_config.json", "Gesture-to-action mappings, available actions, and settings"],
            ["templates/index.html", "Dashboard HTML structure"],
            ["static/css/style.css", "Dark-themed CSS with responsive layout"],
            ["static/js/app.js", "Frontend logic for polling, rendering, and configuration"],
            ["requirements.txt", "Python package dependencies"],
        ],
    )

    # ==========================================================
    # 12. LIMITATIONS
    # ==========================================================
    add_heading(doc, "12. Limitations")

    limitations = [
        ("Lighting sensitivity", "Detection accuracy drops noticeably in very "
         "low or uneven lighting conditions. MediaPipe relies on clear hand "
         "visibility to detect landmarks accurately."),
        ("Single-hand tracking", "The system currently tracks only one hand at "
         "a time. Supporting two hands could enable a richer gesture vocabulary "
         "but would increase processing load."),
        ("Angle sensitivity", "Extreme hand tilts (beyond roughly 60 degrees) "
         "can cause finger misclassification because the vertical tip-above-knuckle "
         "check becomes unreliable at steep angles."),
        ("Active window requirement", "PyAutoGUI sends keystrokes to whichever "
         "window is focused. The user must keep their media player in focus for "
         "the shortcuts to work correctly."),
        ("Camera requirement", "The system requires a working webcam. Desktops "
         "without a built-in camera need an external USB webcam."),
    ]
    for lname, ldesc in limitations:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{lname} — ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Calibri"
        run2 = p.add_run(ldesc)
        run2.font.size = Pt(11)
        run2.font.name = "Calibri"

    # ==========================================================
    # 13. FUTURE WORK
    # ==========================================================
    add_heading(doc, "13. Future Work")

    future_items = [
        "Two-hand gesture support to expand the number of available commands",
        "Custom gesture training using collected landmark data as features for a classifier",
        "Browser extension integration to eliminate the active-window dependency",
        "Volume control via hand height (raise hand to increase, lower to decrease)",
        "Mobile camera input support through IP Webcam or DroidCam integration",
    ]
    for item in future_items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.size = Pt(11)
            run.font.name = "Calibri"

    # ==========================================================
    # 14. CONCLUSION
    # ==========================================================
    add_heading(doc, "14. Conclusion")

    add_para(doc,
        "GesturePlay shows that a practical, touchless media controller can "
        "be built using nothing more than a laptop webcam and a few open-source "
        "libraries. The combination of MediaPipe for hand tracking, an 8-frame "
        "stability buffer for noise filtering, and neutral-gap gating for "
        "transition-state prevention results in a system that works reliably "
        "in everyday conditions."
    )
    add_para(doc,
        "The web dashboard provides a clean interface for monitoring and "
        "customization, and the configurable gesture mappings mean users can "
        "adapt the system to their own preferences. While there are clear "
        "limitations — especially around lighting and viewing angle — the "
        "project demonstrates a viable approach to gesture-based human-computer "
        "interaction that can be extended and improved over time."
    )

    # Save
    doc.save(OUTPUT)
    print(f"Report saved to: {OUTPUT}")


if __name__ == "__main__":
    build_report()
